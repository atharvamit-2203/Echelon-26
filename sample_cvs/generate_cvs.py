"""
Generate PDF and Word format CVs for Fair-Hire Sentinel testing
Run: python generate_cvs.py
"""

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# All CV data
cvs = [
    {
        "filename": "Robert_Chen_CTO",
        "name": "Robert Chen",
        "age": 52,
        "gender": "Male",
        "email": "robert.chen@email.com",
        "phone": "+1-555-0187",
        "location": "San Francisco, CA",
        "summary": "Visionary Chief Technology Officer with 25+ years of experience leading digital transformation initiatives at Fortune 500 companies. Expert in enterprise architecture, cloud migration, and building high-performing engineering teams of 200+ members.",
        "experience": [
            ("Chief Technology Officer", "TechCorp Global", "2018 - Present", [
                "Led digital transformation saving $50M annually",
                "Architected cloud migration for 200+ applications",
                "Built and managed engineering teams across 5 global offices",
                "Implemented AI/ML solutions increasing efficiency by 40%"
            ]),
            ("VP of Engineering", "DataSystems Inc", "2012 - 2018", [
                "Scaled engineering from 30 to 150 engineers",
                "Led architecture redesign improving system performance by 300%",
                "Championed DevOps culture reducing deployment time by 80%"
            ])
        ],
        "skills": "Cloud Architecture, AWS, Azure, Kubernetes, Microservices, System Design, Leadership, Strategy, Agile, DevOps, Python, Java, C++, Machine Learning, Data Architecture, Security, Scalability",
        "education": ["MS Computer Science - Stanford University", "BS Computer Engineering - MIT"]
    },
    # Add more CVs here...
]

def create_word_cv(cv_data):
    """Create a Word document CV"""
    doc = Document()
    
    # Name (Title)
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(cv_data['name'])
    name_run.font.size = Pt(18)
    name_run.font.bold = True
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Contact Info
    contact = doc.add_paragraph()
    contact_text = f"Age: {cv_data['age']} | Gender: {cv_data['gender']} | {cv_data['email']} | {cv_data['phone']}"
    contact.add_run(contact_text)
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(cv_data['location']).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Summary
    doc.add_paragraph()
    summary_heading = doc.add_paragraph()
    summary_heading.add_run('PROFESSIONAL SUMMARY').bold = True
    doc.add_paragraph(cv_data['summary'])
    
    # Experience
    doc.add_paragraph()
    exp_heading = doc.add_paragraph()
    exp_heading.add_run('WORK EXPERIENCE').bold = True
    
    for title, company, period, bullets in cv_data['experience']:
        doc.add_paragraph()
        job = doc.add_paragraph()
        job.add_run(f"{title} | {company} | {period}").bold = True
        for bullet in bullets:
            doc.add_paragraph(f"• {bullet}", style='List Bullet')
    
    # Skills
    doc.add_paragraph()
    skills_heading = doc.add_paragraph()
    skills_heading.add_run('SKILLS').bold = True
    doc.add_paragraph(cv_data['skills'])
    
    # Education
    doc.add_paragraph()
    edu_heading = doc.add_paragraph()
    edu_heading.add_run('EDUCATION').bold = True
    for edu in cv_data['education']:
        doc.add_paragraph(edu)
    
    # Save
    filename = f"{cv_data['filename']}.docx"
    doc.save(filename)
    print(f"Created Word: {filename}")

def create_pdf_cv(cv_data):
    """Create a PDF CV"""
    pdf_path = f"{cv_data['filename']}.pdf"
    doc = SimpleDocTemplate(pdf_path, pagesize=letter,
                            rightMargin=72, leftMargin=72,
                            topMargin=72, bottomMargin=18)
    
    elements = []
    styles = getSampleStyleSheet()
    
    # Name
    elements.append(Paragraph(f"<b>{cv_data['name']}</b>", styles['Title']))
    elements.append(Spacer(1, 0.1*inch))
    
    # Contact
    contact = f"Age: {cv_data['age']} | Gender: {cv_data['gender']}<br/>{cv_data['email']} | {cv_data['phone']}<br/>{cv_data['location']}"
    elements.append(Paragraph(contact, styles['Normal']))
    elements.append(Spacer(1, 0.2*inch))
    
    # Summary
    elements.append(Paragraph("<b>PROFESSIONAL SUMMARY</b>", styles['Heading2']))
    elements.append(Paragraph(cv_data['summary'], styles['Normal']))
    elements.append(Spacer(1, 0.2*inch))
    
    # Experience
    elements.append(Paragraph("<b>WORK EXPERIENCE</b>", styles['Heading2']))
    for title, company, period, bullets in cv_data['experience']:
        elements.append(Spacer(1, 0.1*inch))
        elements.append(Paragraph(f"<b>{title} | {company} | {period}</b>", styles['Normal']))
        for bullet in bullets:
            elements.append(Paragraph(f"• {bullet}", styles['Normal']))
    elements.append(Spacer(1, 0.2*inch))
    
    # Skills
    elements.append(Paragraph("<b>SKILLS</b>", styles['Heading2']))
    elements.append(Paragraph(cv_data['skills'], styles['Normal']))
    elements.append(Spacer(1, 0.2*inch))
    
    # Education
    elements.append(Paragraph("<b>EDUCATION</b>", styles['Heading2']))
    for edu in cv_data['education']:
        elements.append(Paragraph(edu, styles['Normal']))
    
    doc.build(elements)
    print(f"Created PDF: {pdf_path}")

if __name__ == "__main__":
    print("Generating CVs in PDF and Word formats...\n")
    
    for cv in cvs:
        try:
            create_word_cv(cv)
            create_pdf_cv(cv)
        except Exception as e:
            print(f"Error creating CV for {cv['filename']}: {e}")
    
    print(f"\nSuccessfully generated {len(cvs)} CVs in both formats!")
    print("\nNote: You can upload these PDF and Word files to test the bias detection system.")
